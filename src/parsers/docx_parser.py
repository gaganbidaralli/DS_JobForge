"""Extract text from DOCX files."""


def extract_text_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(lines)
