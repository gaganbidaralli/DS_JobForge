"""
docx_generator.py  --  Resume DOCX Exporter  (FIXED v1.2)

BUGS FIXED
----------
1. Curly braces in bullet points
   Same root cause as pdf_generator: responsibilities stored as dicts were
   converted to str(), producing "{'bullet_point': 'text', ...}" in the doc.
   Fix: _text() helper extracts plain string from any bullet format.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helper: safely extract plain text from any bullet format
# ---------------------------------------------------------------------------

def _text(item) -> str:
    """Extract a plain string from str, dict, or any other type."""
    if isinstance(item, str):
        return item.strip()
    if isinstance(item, dict):
        for key in ("bullet_point", "text", "description", "point",
                    "responsibility", "detail", "content"):
            if key in item and isinstance(item[key], str) and item[key].strip():
                return item[key].strip()
        for v in item.values():
            if isinstance(v, str) and v.strip():
                return v.strip()
        return ""
    return str(item).strip()


def _skills_list(val) -> list:
    if isinstance(val, list):
        return [_text(s) for s in val if _text(s)]
    if isinstance(val, str):
        return [s.strip() for s in val.split(",") if s.strip()]
    return []


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _add_hrule(doc: Document) -> None:
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    _add_hrule(doc)


def _job_title_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2c, 0x2c, 0x54)


def _meta_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def _bullet_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(9)


def _body_para(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(9)
    run = p.add_run(text)
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# DOCX Generator
# ---------------------------------------------------------------------------

class DOCXResumeGenerator:

    def generate_resume(self, resume: dict, output_path) -> str:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.75)
            section.right_margin  = Inches(0.75)

        # Default font
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        # ── Header ────────────────────────────────────────────────────────────
        info = resume.get("personal_info", {})
        name = info.get("full_name", "Resume") or "Resume"

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)
        nr = name_para.add_run(name)
        nr.bold = True
        nr.font.size = Pt(20)
        nr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        contact_parts = [
            info.get("email", ""),
            info.get("phone", ""),
            info.get("location", ""),
            info.get("linkedin", ""),
        ]
        contact_line = "  |  ".join(p for p in contact_parts if p)
        if contact_line:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(4)
            cr = cp.add_run(contact_line)
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        _add_hrule(doc)

        # ── Summary ───────────────────────────────────────────────────────────
        summary = resume.get("summary", "")
        if summary and isinstance(summary, str) and summary.strip():
            _section_heading(doc, "Professional Summary")
            _body_para(doc, summary.strip())

        # ── Experience ────────────────────────────────────────────────────────
        experience = resume.get("experience", [])
        if experience:
            _section_heading(doc, "Experience")
            for job in experience:
                if not isinstance(job, dict):
                    continue
                position = job.get("position", "") or ""
                company  = job.get("company", "")  or ""
                title_str = (
                    f"{position} — {company}" if position and company
                    else position or company
                )
                if title_str:
                    _job_title_para(doc, title_str)

                start = job.get("start_date", "") or ""
                end   = job.get("end_date", "")   or ""
                loc   = job.get("location", "")   or ""
                date_str = f"{start} – {end}" if start or end else ""
                meta = "  |  ".join(p for p in [date_str, loc] if p)
                if meta:
                    _meta_para(doc, meta)

                # FIXED: extract plain text from each responsibility
                for resp in job.get("responsibilities", []):
                    txt = _text(resp)
                    if txt:
                        _bullet_para(doc, txt)

        # ── Education ─────────────────────────────────────────────────────────
        education = resume.get("education", [])
        if education:
            _section_heading(doc, "Education")
            for edu in education:
                if not isinstance(edu, dict):
                    continue
                inst   = edu.get("institution", "")     or ""
                degree = edu.get("degree", "")          or ""
                field  = edu.get("field", "")           or ""
                year   = edu.get("graduation_year", "") or ""
                degree_str = f"{degree} in {field}" if degree and field else degree or field
                if inst:
                    _job_title_para(doc, f"{inst}  —  {degree_str}" if degree_str else inst)
                if year:
                    _meta_para(doc, f"Graduated: {year}")

        # ── Skills ────────────────────────────────────────────────────────────
        skills = resume.get("skills", {})
        if skills and isinstance(skills, dict):
            labels = {
                "technical": "Technical",
                "soft":      "Soft Skills",
                "tools":     "Tools",
                "languages": "Languages",
            }
            skill_lines = []
            for key, label in labels.items():
                items = _skills_list(skills.get(key, []))
                if items:
                    skill_lines.append((label, ", ".join(items)))

            if skill_lines:
                _section_heading(doc, "Skills")
                for label, line in skill_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    r1 = p.add_run(f"{label}: ")
                    r1.bold = True
                    r1.font.size = Pt(9)
                    r2 = p.add_run(line)
                    r2.font.size = Pt(9)

        # ── Certifications ────────────────────────────────────────────────────
        certs = resume.get("certifications", [])
        if certs:
            _section_heading(doc, "Certifications")
            for cert in certs:
                txt = _text(cert)
                if txt:
                    _bullet_para(doc, txt)

        # ── Projects ──────────────────────────────────────────────────────────
        projects = resume.get("projects", [])
        if projects:
            _section_heading(doc, "Projects")
            for proj in projects:
                if not isinstance(proj, dict):
                    _bullet_para(doc, _text(proj))
                    continue
                proj_name = proj.get("name", "") or proj.get("title", "") or ""
                if proj_name:
                    _job_title_para(doc, proj_name)
                # Technologies line
                tech = proj.get("technologies", [])
                if isinstance(tech, str):
                    tech = [t.strip() for t in tech.split(",") if t.strip()]
                if tech:
                    _body_para(doc, ", ".join(_text(t) for t in tech if _text(t)), "Tech: ")
                # Description
                desc = proj.get("description", "")
                if desc:
                    _body_para(doc, _text(desc))
                # Highlights / bullets
                for field in ("highlights", "bullets", "responsibilities"):
                    for b in proj.get(field, []):
                        txt = _text(b)
                        if txt:
                            _bullet_para(doc, txt)
                # URL
                url = proj.get("url", "")
                if url and isinstance(url, str) and url.strip():
                    _body_para(doc, url.strip(), "Link: ")

        doc.save(str(output_path))
        return str(output_path)